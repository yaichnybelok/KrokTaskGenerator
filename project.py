import random
import tkinter

import yaml
from PIL import ImageGrab
from docx import *
from docx.shared import Inches, Cm, Pt

with open('config.yaml', 'r', encoding='utf-8') as file:
    configuration = yaml.safe_load(file)

alphabet = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
trajectory_list = []  # cписок для записи траекторий
trajectory_coords = []
answer = []
Document_Name = ''
WINDOW_WIDTH = configuration['WINDOW_WIDTH']
WINDOW_HEIGHT = configuration['WINDOW_HEIGHT']
CELLSIZE = configuration['CELLSIZE']
TRAJECTORY_LENGTH = configuration['TRAJECTORY_LENGTH']
TRAJECTORY_AMOUNT = configuration['TRAJECTORY_AMOUNT']
TRAJECTORY_STEP = configuration['TRAJECTORY_STEP']
TASK_START_POINT_X = configuration['TASK_START_POINT_X']
TASK_START_POINT_Y = configuration['TASK_START_POINT_Y']
DIFFICULTY = configuration['DIFFICULTY']
CELLS = configuration['CELLS']
TASK_DESCRIPTION = configuration['TASK_DESCRIPTION']
TASKS_AMOUNT = configuration['TASKS_AMOUNT']
TASKS_FILENAME = configuration['TASKS_FILENAME']
ANSWERS_FILENAME = configuration['ANSWERS_FILENAME']
WINDOW_COORD_X = 0
WINDOW_COORD_Y = 0


# --------------------------------------------FUNCTIONS-----------------------------------------------------------------
def CellSizing(n):
    return n * CELLSIZE


def TrajectoryRotater(n, a):
    angle = random.choice([0, 90, 180, 270])
    l = ['up', 'right', 'down', 'left', 'up', 'right', 'down', 'left']
    rotated = []

    for i in trajectory_list[a][n]:
        if angle == 0:
            rotated.append(l[l.index(i) + 0])
        elif angle == 90:
            rotated.append(l[l.index(i) + 1])
        elif angle == 180:
            rotated.append(l[l.index(i) + 2])
        elif angle == 270:
            rotated.append(l[l.index(i) + 3])

    return rotated


def MakeField():
    for i in range(TASKS_AMOUNT):
        answer.append([])

    document = Document()

    for i in range(TASKS_AMOUNT):
        trajectory_list.append([])
        trajectory_coords.append([])
        for j in range(TRAJECTORY_AMOUNT):
            trajectory_list[i].append([])
            trajectory_coords[i].append([])

    for a in range(TASKS_AMOUNT):
        document.add_paragraph('Задание "Траектория" ' + str(a + 1))
        document.add_paragraph(TASK_DESCRIPTION)

        master = tkinter.Tk()
        canvas = tkinter.Canvas(master, bg='white', height=WINDOW_HEIGHT, width=WINDOW_WIDTH)  # Cоздание холста
        master.attributes('-fullscreen', False)
        if CELLS == True:
            for i in range(CELLSIZE, WINDOW_WIDTH, CELLSIZE):
                canvas.create_line((i, 0), (i, WINDOW_HEIGHT), fill='black')
            for i in range(CELLSIZE, WINDOW_WIDTH, CELLSIZE):
                canvas.create_line((0, i), (WINDOW_WIDTH, i), fill='black')
        # Рисование клеток

        DefaultPoint = [CellSizing(8), CellSizing(6)]
        TrajectoryStartPoint = [CellSizing(7), CellSizing(6)]  # Исходная точка траектории

        # Генерация траекторий

        for j in range(TRAJECTORY_AMOUNT):
            canvas.create_oval((TrajectoryStartPoint[0], TrajectoryStartPoint[1]),
                               (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=9,
                               outline='black')  # рисуем точку начала траектории

            PreviousRandomNumber = 0

            for i in range(TRAJECTORY_LENGTH):
                trajectory_coords[a][j].append(TrajectoryStartPoint)
                RandomNumber = random.randint(1, 4)  # Случайное число для генерации поворотов
                lenght = 1  # длина поворота

                while abs(PreviousRandomNumber - RandomNumber) == 2:
                    RandomNumber = random.randint(1, 4)

                if RandomNumber == 4:
                    canvas.create_line((TrajectoryStartPoint[0], TrajectoryStartPoint[1] - CellSizing(1)),
                                       (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=3)

                    TrajectoryStartPoint[1] = TrajectoryStartPoint[1] - CellSizing(1)
                    trajectory_list[a][j].append('up')

                elif RandomNumber == 2:
                    canvas.create_line((TrajectoryStartPoint[0], TrajectoryStartPoint[1] + CellSizing(1)),
                                       (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=3)

                    TrajectoryStartPoint[1] = TrajectoryStartPoint[1] + CellSizing(1)
                    trajectory_list[a][j].append('down')

                else:
                    match RandomNumber:
                        case 3:
                            step = CellSizing(-lenght)
                            trajectory_list[a][j].append('left')
                        case 1:
                            step = CellSizing(lenght)
                            trajectory_list[a][j].append('right')
                    canvas.create_line((TrajectoryStartPoint[0], TrajectoryStartPoint[1]),
                                       (TrajectoryStartPoint[0] + step, TrajectoryStartPoint[1]), fill='black',
                                       width=3)  # рисуем линию
                    TrajectoryStartPoint[0] = TrajectoryStartPoint[0] + step

                PreviousRandomNumber = RandomNumber
                canvas.create_oval((TrajectoryStartPoint[0], TrajectoryStartPoint[1]),
                                   (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=5,
                                   outline='black')
            TrajectoryStartPoint = [DefaultPoint[0] + CellSizing((j + 1) * TRAJECTORY_STEP), DefaultPoint[1]]

        # ---------------------------------------------Генерация задачи-------------------------------------------------

        TrajectoryStartPoint = [CellSizing(TASK_START_POINT_X), CellSizing(TASK_START_POINT_Y)]
        canvas.create_oval((TrajectoryStartPoint[0], TrajectoryStartPoint[1]),
                           (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=9,
                           outline='black')

        for n in range(DIFFICULTY):
            trajectory_num = random.randint(0, TRAJECTORY_AMOUNT - 1)
            for i in TrajectoryRotater(trajectory_num, a):
                if i == 'up':
                    canvas.create_line((TrajectoryStartPoint[0], TrajectoryStartPoint[1] - CellSizing(1)),
                                       (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=3)

                    TrajectoryStartPoint[1] = TrajectoryStartPoint[1] - CellSizing(1)


                elif i == 'right':
                    canvas.create_line((TrajectoryStartPoint[0] + CellSizing(1)), TrajectoryStartPoint[1],
                                       (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=3)

                    TrajectoryStartPoint[0] = TrajectoryStartPoint[0] + CellSizing(1)

                elif i == 'down':
                    canvas.create_line((TrajectoryStartPoint[0], TrajectoryStartPoint[1] + CellSizing(1)),
                                       (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=3)
                    TrajectoryStartPoint[1] = TrajectoryStartPoint[1] + CellSizing(1)

                elif i == 'left':
                    canvas.create_line((TrajectoryStartPoint[0] - CellSizing(1)), TrajectoryStartPoint[1],
                                       (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=3)

                    TrajectoryStartPoint[0] = TrajectoryStartPoint[0] - CellSizing(1)
                canvas.create_oval((TrajectoryStartPoint[0], TrajectoryStartPoint[1]),
                                   (TrajectoryStartPoint[0], TrajectoryStartPoint[1]), fill='black', width=5,
                                   outline='black')
            answer[a].append(alphabet[trajectory_num])

        # Сохранение файла

        label = tkinter.Label(master, text="Генератор задач с роботом", relief='raised')
        label.pack()
        canvas.pack()
        master.mainloop()
        screenshot = ImageGrab.grab(
            bbox=(15, 70, 1000, 1000))
        screenshot.save('screenshot.jpg')
        document.add_picture('screenshot.jpg', width=Cm(15), height=Cm(15))
        document.add_paragraph(' ')
        document.add_paragraph(' ')
        document.add_paragraph('Ответ:____________')
        document.save(str(TASKS_FILENAME) + '.docx')
        ans_document = Document()
        for i in range(len(answer)):
            ans_document.add_paragraph('Задание ' + str(i + 1) + ', ответ:')
            answer_str = ''
            for j in answer[i]:
                answer_str += j
            ans_document.add_paragraph(answer_str)
            ans_document.save(str(ANSWERS_FILENAME) + '.docx')


# ----------------------------------------------MAIN--------------------------------------------------------------------


MakeField()
